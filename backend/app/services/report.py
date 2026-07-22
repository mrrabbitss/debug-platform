from pathlib import Path
from typing import Any

from docx import Document
from jinja2 import Environment, FileSystemLoader, select_autoescape
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy import func, select

from app.core.config import get_settings
from app.core.db import SessionLocal
from app.core.utils import json_loads, new_id, sha256_file
from app.models import AnalysisRun, Case, Report
from app.services.storage import storage


def get_report_context(case_id: str, analysis_id: str) -> dict[str, Any]:
    with SessionLocal() as db:
        case = db.get(Case, case_id)
        analysis = db.get(AnalysisRun, analysis_id)
        if not case or not analysis or analysis.case_id != case_id:
            raise ValueError("Case or analysis not found")
        result = json_loads(analysis.result_json, {})
        evidence = {item.get("evidence_id"): item for item in json_loads(analysis.evidence_json, [])}
    return {
        "title": get_settings().report_title,
        "case": case,
        "analysis": analysis,
        "result": result,
        "evidence": evidence,
    }


def render_html(case_id: str, analysis_id: str) -> str:
    context = get_report_context(case_id, analysis_id)
    template_dir = Path(__file__).resolve().parents[1] / "templates"
    env = Environment(loader=FileSystemLoader(template_dir), autoescape=select_autoescape(["html", "xml"]))
    return env.get_template("report.html.j2").render(**context)


def _next_version(case_id: str, analysis_id: str, fmt: str) -> int:
    with SessionLocal() as db:
        current = db.scalar(select(func.max(Report.version)).where(
            Report.case_id == case_id, Report.analysis_run_id == analysis_id, Report.format == fmt
        ))
    return int(current or 0) + 1


def _record_report(case_id: str, analysis_id: str, fmt: str, path: Path, version: int) -> Report:
    report = Report(
        id=new_id("RPT"), case_id=case_id, analysis_run_id=analysis_id,
        format=fmt, version=version, stored_path=storage.storage_key(path), sha256=sha256_file(path),
    )
    with SessionLocal() as db:
        db.add(report)
        db.commit()
        db.refresh(report)
    return report


def generate_html_file(case_id: str, analysis_id: str) -> Report:
    version = _next_version(case_id, analysis_id, "html")
    path = storage.report_dir(case_id) / f"{analysis_id}_v{version}.html"
    path.write_text(render_html(case_id, analysis_id), encoding="utf-8")
    return _record_report(case_id, analysis_id, "html", path, version)


def generate_docx(case_id: str, analysis_id: str) -> Report:
    context = get_report_context(case_id, analysis_id)
    result = context["result"]
    case = context["case"]
    version = _next_version(case_id, analysis_id, "docx")
    path = storage.report_dir(case_id) / f"{analysis_id}_v{version}.docx"
    document = Document()
    document.add_heading(context["title"], 0)
    document.add_heading("一、基本信息", level=1)
    for label, value in [
        ("案例编号", case.id), ("问题标题", case.title), ("设备类型", case.device_type),
        ("设备型号", case.device_model or "未提供"), ("固件版本", case.firmware_version or "未提供"),
        ("问题发生时间", case.issue_time or "未提供"),
    ]:
        document.add_paragraph(f"{label}：{value}")
    document.add_heading("二、综合摘要", level=1)
    document.add_paragraph(result.get("summary", "暂无"))
    document.add_heading("三、已确认事实", level=1)
    for fact in result.get("confirmed_facts", []):
        document.add_paragraph(fact.get("statement", ""), style="List Bullet")
    document.add_heading("四、根因候选", level=1)
    for item in result.get("hypotheses", []):
        document.add_heading(f"{item.get('rank', '-')}. {item.get('title', '')}", level=2)
        document.add_paragraph(item.get("description", ""))
        document.add_paragraph(f"可信等级：{item.get('confidence_level', 'UNKNOWN')}；优先级：{item.get('priority', 'UNKNOWN')}")
        document.add_paragraph("证据：" + "、".join(item.get("supporting_evidence", [])))
    document.add_heading("五、建议排查步骤", level=1)
    for action in result.get("recommended_actions", []):
        document.add_paragraph(f"[{action.get('priority')}] {action.get('action')} — {action.get('reason')}", style="List Number")
    document.add_heading("六、缺失信息与限制", level=1)
    for item in result.get("missing_information", []) + result.get("limitations", []):
        document.add_paragraph(item, style="List Bullet")
    document.save(path)
    return _record_report(case_id, analysis_id, "docx", path, version)


def generate_pdf(case_id: str, analysis_id: str) -> Report:
    context = get_report_context(case_id, analysis_id)
    result = context["result"]
    case = context["case"]
    version = _next_version(case_id, analysis_id, "pdf")
    path = storage.report_dir(case_id) / f"{analysis_id}_v{version}.pdf"
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("CNTitle", parent=styles["Title"], alignment=TA_CENTER, fontName="STSong-Light")
    heading = ParagraphStyle("CNHeading", parent=styles["Heading2"], fontName="STSong-Light", spaceBefore=8)
    body = ParagraphStyle("CNBody", parent=styles["BodyText"], fontName="STSong-Light", leading=15)
    story = [Paragraph("GW/AP Intelligent Diagnosis Report", title_style), Spacer(1, 6 * mm)]
    info = [
        ["Case ID", case.id], ["Title", case.title], ["Device", f"{case.device_type} / {case.device_model or '-'}"],
        ["Firmware", case.firmware_version or "-"], ["Issue time", case.issue_time or "-"],
    ]
    table = Table(info, colWidths=[35 * mm, 140 * mm])
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.grey), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story += [table, Spacer(1, 5 * mm), Paragraph("Summary", heading), Paragraph(result.get("summary", "N/A"), body)]
    story.append(Paragraph("Confirmed facts", heading))
    for fact in result.get("confirmed_facts", [])[:30]:
        story.append(Paragraph("• " + fact.get("statement", ""), body))
    story.append(Paragraph("Root-cause hypotheses", heading))
    for item in result.get("hypotheses", []):
        story.append(Paragraph(f"{item.get('rank')}. {item.get('title')} [{item.get('confidence_level')}]", body))
        story.append(Paragraph(item.get("description", ""), body))
    story.append(PageBreak())
    story.append(Paragraph("Recommended actions", heading))
    for action in result.get("recommended_actions", []):
        story.append(Paragraph(f"[{action.get('priority')}] {action.get('action')} — {action.get('reason')}", body))
    SimpleDocTemplate(str(path), pagesize=A4, rightMargin=15 * mm, leftMargin=15 * mm, topMargin=15 * mm, bottomMargin=15 * mm).build(story)
    return _record_report(case_id, analysis_id, "pdf", path, version)
